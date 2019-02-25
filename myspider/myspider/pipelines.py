# -*- coding: utf-8 -*-

# Define your item pipelines here
#
# Don't forget to add your pipeline to the ITEM_PIPELINES setting
# See: https://doc.scrapy.org/en/latest/topics/item-pipeline.html
import docx
from . import Global
from docx import *


class MyspiderPipeline(object):
    document = Document()

    def __init__(self):
        self.document = Document()

    def process_item(self, item, spider):
        para = self.document.add_paragraph(item.runid)
        para = self.document.add_paragraph(item.title)
        para = self.document.add_paragraph(item.date)
        return item

    def __del__(self):
        self.document.save(Global.Glovar.stuid + '.docx')
